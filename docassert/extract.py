"""Extract plain text from a source document, for doc-to-pmo conversion.

The deterministic first step of the conversion front-door: turn an arbitrary
source file (.docx / .pdf / .md / .txt) into plain text that the doc-to-pmo
skill then maps into a standard template. It does not interpret or reshape the
content — that is the skill's job.

.docx / .pdf support needs the optional `convert` extra:
    pip install "docassert[convert]"
"""
from __future__ import annotations

from pathlib import Path

_NEED_CONVERT = 'extract needs the "convert" extra: pip install "docassert[convert]"'


def extract(path: str | Path) -> str:
    """Return the plain text of a source document.

    Raises FileNotFoundError (missing file), ValueError (unsupported type), or
    ImportError (a .docx/.pdf without the `convert` extra installed).
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"no such file: {p}")
    ext = p.suffix.lower()

    if ext in {".md", ".txt"}:
        return p.read_text(encoding="utf-8")

    if ext == ".docx":
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(_NEED_CONVERT) from exc
        document = docx.Document(str(p))
        blocks: list[str] = [para.text for para in document.paragraphs]
        # include table cell text, which charters often use for milestones/risks
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ImportError(_NEED_CONVERT) from exc
        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    raise ValueError(f"unsupported source type '{ext}' (supported: .docx, .pdf, .md, .txt)")
